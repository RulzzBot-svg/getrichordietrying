#!/usr/bin/env python3
"""
generate_qr_labels.py — Universal QR code generator.

Phase 5 of the white-label roadmap: QR codes now point to a generic
scan route so the same physical label works for ANY asset type and
ANY industry.

QR URL format:
    <base_url>/scan?asset_id=<asset_db_id>

When the React app opens that URL, it fetches /api/scan?asset_id=<id>,
which returns the asset type and required form fields.  The UI then
renders the correct checklist dynamically — no hardcoded logic.

Usage examples
--------------
  # All assets for a tenant
  python generate_qr_labels.py --tenant acme-hvac --out qr_codes

  # Specific location within a tenant
  python generate_qr_labels.py --tenant acme-hvac --location "Main Campus" --out qr_codes

  # Override the app base URL (e.g. production domain)
  python generate_qr_labels.py --tenant acme-hvac --base-url https://app.example.com

  # Also produce a Word doc with embedded QR images
  python generate_qr_labels.py --tenant acme-hvac --doc

Legacy option (backward compat):
  python generate_qr_labels.py --hospital "Foothill" --out qr_codes
"""
import os
import argparse
import qrcode
from app import create_app
from db import db
from models import Asset, Location, Tenant

try:
    from docx import Document
    from docx.shared import Inches
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def ensure_dir(path: str):
    os.makedirs(path, exist_ok=True)


def make_qr_image(url: str, out_path: str, box_size: int = 6):
    qr = qrcode.QRCode(
        version=2,
        error_correction=qrcode.constants.ERROR_CORRECT_Q,
        box_size=box_size,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    img.save(out_path)


def _scan_url(base_url: str, asset_id: int) -> str:
    """
    Build the canonical scan URL for an asset.
    Format: <base_url>/scan?asset_id=<asset_id>
    """
    return f"{base_url.rstrip('/')}/scan?asset_id={asset_id}"


# ---------------------------------------------------------------------------
# Core generator
# ---------------------------------------------------------------------------
def gen_for_location(location: Location, out_dir: str, base_url: str, make_doc: bool = False):
    """Generate QR PNGs (and optional Word doc) for all assets at a location."""
    ensure_dir(out_dir)

    assets = Asset.query.filter_by(location_id=location.id).order_by(Asset.id).all()
    if not assets:
        print(f"  No assets found for location: {location.name}")
        return

    doc = None
    if make_doc:
        if not HAVE_DOCX:
            print("  python-docx not available; skipping Word document creation")
            make_doc = False
        else:
            doc = Document()
            doc.add_heading(f"QR Codes — {location.name}", 0)

    for asset in assets:
        url = _scan_url(base_url, asset.id)
        filename = f"{asset.id}.png"
        out_path = os.path.join(out_dir, filename)
        make_qr_image(url, out_path)
        print(f"  ✓ {out_path}  →  {url}")

        if make_doc and doc is not None:
            doc.add_heading(f"{asset.name}", level=1)
            if asset.asset_type:
                doc.add_paragraph(f"Type: {asset.asset_type}")
            if asset.location_label:
                doc.add_paragraph(f"Location: {asset.location_label}")
            try:
                doc.add_picture(out_path, width=Inches(2))
            except Exception as e:
                print(f"  ⚠ Failed to embed picture in doc: {e}")
            doc.add_paragraph("")

    if make_doc and doc is not None:
        safe_name = location.name.replace(" ", "_").replace("/", "-")
        doc_path = os.path.join(out_dir, f"{safe_name}_QRs.docx")
        doc.save(doc_path)
        print(f"  📄 Word doc saved: {doc_path}")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def build_parser():
    p = argparse.ArgumentParser(
        description="Generate QR code PNGs for assets.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    # Tenant / target selection
    p.add_argument("--tenant", "-t", help="Tenant slug (generates for all locations in that tenant)")
    p.add_argument("--location", "-l", help="Location name within the tenant (optional filter)")

    # Legacy HVAC backward-compat flags
    p.add_argument("--hospital", "-H", help="[Legacy] Hospital name (same as --location, no tenant filter)")
    p.add_argument("--all", action="store_true", help="[Legacy] Generate for all locations/hospitals")

    p.add_argument("--out", "-o", default="qr_codes", help="Output directory (default: qr_codes)")
    p.add_argument(
        "--base-url", "-b",
        default="https://qrscan-lyart.vercel.app",
        help="App base URL embedded in QR codes (default: https://qrscan-lyart.vercel.app)",
    )
    p.add_argument("--doc", action="store_true", help="Also produce a Word document with embedded images")
    return p


def main():
    parser = build_parser()
    args = parser.parse_args()

    app = create_app()
    with app.app_context():
        ensure_dir(args.out)

        # ------------------------------------------------------------------
        # Tenant-aware mode (white-label)
        # ------------------------------------------------------------------
        if args.tenant:
            tenant = Tenant.query.filter_by(slug=args.tenant).first()
            if not tenant:
                print(f"Tenant '{args.tenant}' not found. Available tenants:")
                for t in Tenant.query.order_by(Tenant.slug).all():
                    print(f"  - {t.slug}  ({t.name})")
                return

            locations = Location.query.filter_by(tenant_id=tenant.id).order_by(Location.name).all()

            if args.location:
                locations = [loc for loc in locations if loc.name == args.location]
                if not locations:
                    print(f"Location '{args.location}' not found for tenant '{args.tenant}'.")
                    return

            print(f"Tenant: {tenant.name}  ({len(locations)} location(s))")
            for loc in locations:
                out_dir = os.path.join(args.out, loc.name.replace(" ", "_"))
                print(f"\nLocation: {loc.name}")
                gen_for_location(loc, out_dir, args.base_url, make_doc=args.doc)

        # ------------------------------------------------------------------
        # Legacy mode (--all or --hospital)
        # ------------------------------------------------------------------
        elif args.all or args.hospital:
            if args.all:
                locations = Location.query.order_by(Location.name).all()
                for loc in locations:
                    out_dir = os.path.join(args.out, loc.name.replace(" ", "_"))
                    print(f"\nLocation: {loc.name}")
                    gen_for_location(loc, out_dir, args.base_url, make_doc=args.doc)
            else:
                loc = Location.query.filter_by(name=args.hospital).first()
                if not loc:
                    print(f"Location '{args.hospital}' not found. Available:")
                    for ll in Location.query.order_by(Location.name).all():
                        print(f"  - {ll.name}")
                    return
                gen_for_location(loc, args.out, args.base_url, make_doc=args.doc)

        else:
            parser.print_help()


if __name__ == "__main__":
    main()
